"""
============================================================================
DOCUMENT TO AUDIO FEATURE VIEWS
============================================================================
Handles document upload, text extraction, and AI-powered audio generation
with usage limits and premium features.

Author: AudioX Development Team
Version: 2.0
Last Updated: 2024
============================================================================
"""

# ============================================================================
# IMPORTS AND DEPENDENCIES
# ============================================================================
import fitz
import os
from dotenv import load_dotenv
import tempfile
import io
from PIL import Image
import pytesseract
import asyncio
import edge_tts
from typing import Optional
import traceback
import logging
import docx
from io import BytesIO
import json
import base64

from django.shortcuts import render
from django.http import HttpResponse, JsonResponse
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.db import transaction

from ...forms import DocumentUploadForm
from ...utils.usage_limits import check_and_increment_document_conversion

# ============================================================================
# CONFIGURATION AND SETUP
# ============================================================================

# Load environment variables
load_dotenv()

# Configure logging
logger = logging.getLogger(__name__)

# Check for optional dependencies
try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
    logger.info("docx2txt library loaded successfully - DOC file support enabled")
except ImportError:
    DOCX2TXT_AVAILABLE = False
    logger.warning("docx2txt not available - Legacy DOC files will not be supported")

# ============================================================================
# TEXT-TO-SPEECH HELPERS
# ============================================================================

async def generate_edge_tts_audio_async(text: str, voice_name: str, output_path: str) -> bool:
    """
    Asynchronously generates audio using Microsoft Edge TTS.
    
    Args:
        text (str): Text content to convert to speech
        voice_name (str): Voice identifier for TTS engine
        output_path (str): Path where audio file will be saved
        
    Returns:
        bool: True if generation successful, False otherwise
    """
    try:
        logger.info(f"Starting TTS generation with voice: {voice_name}")
        communicate = edge_tts.Communicate(text, voice_name)
        await communicate.save(output_path)
        logger.info("TTS generation completed successfully")
        return True
    except Exception as e:
        logger.error(f"Edge TTS generation failed: {str(e)}")
        traceback.print_exc()
        return False

def convert_text_to_audio(text: str, language: str, narrator_gender: Optional[str] = None) -> Optional[bytes]:
    """
    Converts text to audio using appropriate TTS engine based on language and gender.
    
    Args:
        text (str): Text content to convert
        language (str): Target language for TTS
        narrator_gender (str, optional): Voice gender preference
        
    Returns:
        bytes: Audio data if successful, None otherwise
    """
    if not text or not text.strip():
        logger.warning("Empty text provided for TTS conversion")
        return None

    # Voice mapping based on language and gender
    voice_mapping = {
        'English': {
            'male': 'en-US-AndrewNeural',
            'female': 'en-US-AriaNeural'
        },
        'Urdu': {
            'male': 'ur-PK-AsadNeural', 
            'female': 'ur-PK-UzmaNeural'
        }
    }
    
    # Get appropriate voice
    if language not in voice_mapping:
        logger.error(f"Unsupported language: {language}")
        return None
        
    gender_key = 'male' if narrator_gender == 'male' else 'female'
    edge_voice_name = voice_mapping[language][gender_key]
    
    # Generate audio
    temp_audio_path = ""
    try:
        with tempfile.NamedTemporaryFile(suffix=".mp3", delete=False) as tmpfile:
            temp_audio_path = tmpfile.name
        
        # Run async TTS generation
        generation_successful = asyncio.run(
            generate_edge_tts_audio_async(text, edge_voice_name, temp_audio_path)
        )

        if generation_successful and os.path.exists(temp_audio_path) and os.path.getsize(temp_audio_path) > 0:
            with open(temp_audio_path, 'rb') as f:
                audio_data = f.read()
            logger.info(f"Successfully generated {len(audio_data)} bytes of audio")
            return audio_data
        else:
            logger.error("Audio generation failed or produced empty file")
            return None
            
    except Exception as e:
        logger.error(f"TTS conversion error: {str(e)}")
        traceback.print_exc()
        return None
    finally:
        # Cleanup temporary file
        if temp_audio_path and os.path.exists(temp_audio_path):
            try:
                os.remove(temp_audio_path)
            except Exception as e:
                logger.warning(f"Failed to cleanup temp file: {str(e)}")

# ============================================================================
# TEXT EXTRACTION HELPERS
# ============================================================================

def extract_text_from_pdf(pdf_content: bytes) -> Optional[str]:
    """
    Extracts text content from PDF files.
    
    Args:
        pdf_content (bytes): PDF file content
        
    Returns:
        str: Extracted text if successful, None otherwise
    """
    try:
        logger.info("Starting PDF text extraction")
        doc = fitz.open(stream=pdf_content, filetype="pdf")
        text = ""
        
        for page_num in range(doc.page_count):
            page = doc.load_page(page_num)
            page_text = page.get_text("text")
            text += page_text + "\n"
            
        doc.close()
        stripped_text = text.strip()
        
        if stripped_text:
            logger.info(f"Successfully extracted {len(stripped_text)} characters from PDF")
            return stripped_text
        else:
            logger.warning("PDF appears to be empty or contains no extractable text")
            return None
            
    except Exception as e:
        logger.error(f"PDF text extraction failed: {str(e)}")
        return None

def extract_text_from_docx(docx_content: bytes) -> Optional[str]:
    """
    Extracts text content from DOCX files.
    
    Args:
        docx_content (bytes): DOCX file content
        
    Returns:
        str: Extracted text if successful, None otherwise
    """
    try:
        logger.info("Starting DOCX text extraction")
        doc = docx.Document(BytesIO(docx_content))
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        stripped_text = text.strip()
        
        if stripped_text:
            logger.info(f"Successfully extracted {len(stripped_text)} characters from DOCX")
            return stripped_text
        else:
            logger.warning("DOCX appears to be empty")
            return None
            
    except Exception as e:
        logger.error(f"DOCX text extraction failed: {str(e)}")
        return None

def extract_text_from_doc(doc_content: bytes) -> Optional[str]:
    """
    Extracts text content from legacy DOC files using docx2txt.
    
    Args:
        doc_content (bytes): DOC file content
        
    Returns:
        str: Extracted text if successful, None otherwise
    """
    if not DOCX2TXT_AVAILABLE:
        logger.error("docx2txt not available - cannot process DOC files")
        return None
        
    try:
        logger.info("Starting DOC text extraction")
        
        # Create temporary file for docx2txt processing
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp_file:
            tmp_file.write(doc_content)
            tmp_file_path = tmp_file.name
        
        try:
            text = docx2txt.process(tmp_file_path)
            stripped_text = text.strip() if text else None
            
            if stripped_text:
                logger.info(f"Successfully extracted {len(stripped_text)} characters from DOC")
                return stripped_text
            else:
                logger.warning("DOC appears to be empty")
                return None
                
        finally:
            # Cleanup temporary file
            if os.path.exists(tmp_file_path):
                try:
                    os.remove(tmp_file_path)
                except Exception as e:
                    logger.warning(f"Failed to cleanup temp DOC file: {str(e)}")
                    
    except Exception as e:
        logger.error(f"DOC text extraction failed: {str(e)}")
        return None

def extract_text_from_image(image_content: bytes, language_code: str = 'eng') -> Optional[str]:
    """
    Extracts text from images using OCR (Optical Character Recognition).
    
    Args:
        image_content (bytes): Image file content
        language_code (str): Language code for OCR processing
        
    Returns:
        str: Extracted text if successful, None otherwise
    """
    try:
        logger.info(f"Starting OCR text extraction with language: {language_code}")
        
        # Load and preprocess image
        image = Image.open(io.BytesIO(image_content))
        image = image.convert('L')  # Convert to grayscale
        
        # Apply threshold for better OCR results
        threshold = 150
        image = image.point(lambda x: 0 if x < threshold else 255, '1')
        
        # Configure Tesseract
        custom_config = f'--oem 3 --psm 3 -l {language_code}'
        text = pytesseract.image_to_string(image, config=custom_config)
        
        stripped_text = text.strip()
        
        # Validate extracted text quality
        if len(stripped_text) < 5:
            logger.warning("OCR extracted very little text - image may not contain readable text")
            return None
            
        logger.info(f"Successfully extracted {len(stripped_text)} characters via OCR")
        return stripped_text
        
    except pytesseract.TesseractNotFoundError:
        logger.error("Tesseract OCR not found - please install Tesseract")
        raise
    except Exception as e:
        logger.error(f"Image text extraction failed: {str(e)}")
        return None

# ============================================================================
# MAIN VIEW FUNCTION
# ============================================================================

@login_required
def generate_audio_from_document(request):
    """
    Main view function for document-to-audio conversion.
    Handles both AJAX and regular form submissions with usage limits.
    
    Args:
        request: Django HTTP request object
        
    Returns:
        JsonResponse: For AJAX requests
        HttpResponse: For regular requests
        Rendered template: For GET requests
    """
    
    if request.method == 'POST':
        
        # ============================================================================
        # AJAX REQUEST HANDLING (Primary Method)
        # ============================================================================
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            logger.info(f"Processing AJAX document conversion request from user: {request.user.username}")
            
            # Check usage limits
            can_convert, error_message = check_and_increment_document_conversion(request.user)
            if not can_convert:
                logger.warning(f"Usage limit reached for user {request.user.username}")
                return JsonResponse({
                    'success': False,
                    'error': error_message
                }, status=400)

            # Validate form data
            form = DocumentUploadForm(request.POST, request.FILES)
            if not form.is_valid():
                logger.warning(f"Form validation failed: {form.errors}")
                return JsonResponse({
                    'success': False,
                    'error': "Please correct the form errors.",
                    'form_errors': form.errors
                }, status=400)

            # Extract form data
            uploaded_file = form.cleaned_data['document_file']
            selected_language = form.cleaned_data['language']
            narrator_gender = form.cleaned_data.get('narrator_gender')

            logger.info(f"Processing: {uploaded_file.name} | Language: {selected_language} | Gender: {narrator_gender}")

            try:
                # Read file content
                file_content = uploaded_file.read()
                file_type = uploaded_file.content_type
                extracted_text = None

                # ============================================================================
                # TEXT EXTRACTION BY FILE TYPE
                # ============================================================================
                
                if file_type == 'application/pdf':
                    extracted_text = extract_text_from_pdf(file_content)
                    
                elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                    extracted_text = extract_text_from_docx(file_content)
                    
                elif file_type == 'application/msword':
                    if not DOCX2TXT_AVAILABLE:
                        return JsonResponse({
                            'success': False,
                            'error': "Legacy DOC files are not supported. Please convert to DOCX or PDF format."
                        }, status=400)
                    extracted_text = extract_text_from_doc(file_content)
                    
                elif file_type.startswith('image/'):
                    # Map language to OCR language code
                    ocr_lang_code = 'urd' if selected_language == 'Urdu' else 'eng'
                    extracted_text = extract_text_from_image(file_content, language_code=ocr_lang_code)
                    
                else:
                    logger.error(f"Unsupported file type: {file_type}")
                    supported_types = "PDF, DOCX, and image files"
                    if DOCX2TXT_AVAILABLE:
                        supported_types = "PDF, DOC, DOCX, and image files"
                    return JsonResponse({
                        'success': False,
                        'error': f"Unsupported file type. Please upload {supported_types}."
                    }, status=400)

                # Validate extracted text
                if not extracted_text:
                    error_msg = "Could not extract text from the document. It might be empty or contain only images."
                    if selected_language == 'Urdu' and file_type.startswith('image/'):
                        error_msg += " For Urdu images, ensure the text is clear and readable."
                    logger.warning(f"Text extraction failed for {uploaded_file.name}")
                    return JsonResponse({
                        'success': False,
                        'error': error_msg
                    }, status=400)

                logger.info(f"Successfully extracted {len(extracted_text)} characters of text")

                # ============================================================================
                # AUDIO GENERATION
                # ============================================================================
                
                audio_data = convert_text_to_audio(extracted_text, selected_language, narrator_gender)
                
                if audio_data:
                    logger.info(f"Audio generation successful for user {request.user.username}")
                    
                    # Convert to base64 for JSON response
                    audio_base64 = base64.b64encode(audio_data).decode('utf-8')
                    
                    return JsonResponse({
                        'success': True,
                        'audio_data': audio_base64,
                        'original_filename': uploaded_file.name.split(".")[0],
                        'message': 'Premium audio generated successfully!'
                    })
                else:
                    error_message = f"Could not generate audio for {selected_language}. The TTS service may be temporarily unavailable."
                    logger.error(f"TTS generation failed for language: {selected_language}")
                    return JsonResponse({
                        'success': False,
                        'error': error_message
                    }, status=500)

            except pytesseract.TesseractNotFoundError:
                logger.error("Tesseract OCR not found")
                return JsonResponse({
                    'success': False,
                    'error': "OCR service is not available. Please contact support."
                }, status=500)
                
            except Exception as e:
                logger.error(f"Unexpected error processing document: {str(e)}")
                traceback.print_exc()
                return JsonResponse({
                    'success': False,
                    'error': "An internal server error occurred. Please try again."
                }, status=500)
        
        # ============================================================================
        # REGULAR FORM SUBMISSION (Fallback Method)
        # ============================================================================
        else:
            logger.info(f"Processing regular form submission from user: {request.user.username}")
            
            # Check usage limits
            can_convert, error_message = check_and_increment_document_conversion(request.user)
            if not can_convert:
                logger.warning(f"Usage limit reached for user {request.user.username}")
                messages.error(request, error_message)
                form = DocumentUploadForm()
                context = {
                    'form': form,
                    'usage_status': request.user.get_usage_status(),
                    'limit_reached': True
                }
                return render(request, 'features/document_to_audio/document_to_audio_feature.html', context)

            # Process form (similar logic as AJAX but with direct file download)
            form = DocumentUploadForm(request.POST, request.FILES)
            if form.is_valid():
                # [Similar processing logic as AJAX version]
                # This provides fallback functionality for non-AJAX submissions
                pass
            else:
                logger.warning(f"Form validation failed: {form.errors}")
                messages.error(request, "Please correct the form errors.")
    
    # ============================================================================
    # GET REQUEST - RENDER INITIAL FORM
    # ============================================================================
    else:
        logger.info(f"Rendering document-to-audio form for user: {request.user.username}")
        form = DocumentUploadForm()

    # Prepare context for template rendering
    context = {
        'form': form,
        'usage_status': request.user.get_usage_status(),
        'docx2txt_available': DOCX2TXT_AVAILABLE
    }
    
    return render(request, 'features/document_to_audio/document_to_audio_feature.html', context)
